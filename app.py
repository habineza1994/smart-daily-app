import os
import datetime
import pymysql
from flask import Flask, request, redirect, session, send_file

app = Flask(__name__)
app.secret_key = "hirwa_secret_key"


# ================= DB =================
def get_db():
    return pymysql.connect(
        host=os.environ.get('MYSQLHOST'),
        user=os.environ.get('MYSQLUSER'),
        password=os.environ.get('MYSQLPASSWORD'),
        database=os.environ.get('MYSQLDATABASE'),
        cursorclass=pymysql.cursors.DictCursor
    )


# ================= INIT DB =================
@app.route("/initdb")
def init_db():
    db = get_db()
    cur = db.cursor()

    cur.execute("""CREATE TABLE IF NOT EXISTS users(
        id INT AUTO_INCREMENT PRIMARY KEY,
        username VARCHAR(100),
        password VARCHAR(255)
    )""")

    cur.execute("""CREATE TABLE IF NOT EXISTS income(
        id INT AUTO_INCREMENT PRIMARY KEY,
        amount DECIMAL(10,2),
        source VARCHAR(255),
        date DATE,
        note TEXT,
        user_id INT
    )""")

    cur.execute("""CREATE TABLE IF NOT EXISTS expenses(
        id INT AUTO_INCREMENT PRIMARY KEY,
        amount DECIMAL(10,2),
        category VARCHAR(255),
        date DATE,
        note TEXT,
        user_id INT
    )""")

    cur.execute("""CREATE TABLE IF NOT EXISTS activities(
        id INT AUTO_INCREMENT PRIMARY KEY,
        activity_name VARCHAR(255),
        done_by VARCHAR(255),
        date DATE,
        description TEXT,
        user_id INT
    )""")

    db.commit()
    return "DB READY"


# ================= AUTH =================
@app.route('/register', methods=['GET','POST'])
def register():
    if request.method == 'POST':
        db = get_db()
        cur = db.cursor()
        cur.execute(
            "INSERT INTO users(username,password) VALUES(%s,%s)",
            (request.form['username'], request.form['password'])
        )
        db.commit()
        return redirect('/login')

    return """
    <h2>Register</h2>
    <form method="POST">
    Username:<input name="username"><br>
    Password:<input name="password"><br>
    <button>Register</button>
    </form>
    """


@app.route("/login", methods=["GET","POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        # logic yawe yo kugenzura login isigare hano

        return redirect("/dashboard")

    return """
<!DOCTYPE html>
<html>
<head>
<title>HIRWA SMART Login</title>
<meta name="viewport" content="width=device-width, initial-scale=1">

<style>
body{
    margin:0;
    font-family: Arial, Helvetica, sans-serif;
    background:linear-gradient(120deg,#4e54c8,#8f94fb);
    height:100vh;
    display:flex;
    justify-content:center;
    align-items:center;
}

.card{
    background:white;
    width:92%;
    max-width:420px;
    border-radius:20px;
    padding:25px;
    box-shadow:0 10px 25px rgba(0,0,0,0.15);
}

.title{
    text-align:center;
    font-size:26px;
    font-weight:bold;
    color:#4e54c8;
    margin-bottom:10px;
}

.subtitle{
    text-align:center;
    color:gray;
    margin-bottom:25px;
}

input{
    width:100%;
    padding:14px;
    margin:10px 0;
    border-radius:10px;
    border:1px solid #ddd;
    font-size:16px;
}

button{
    width:100%;
    padding:14px;
    margin-top:15px;
    border:none;
    border-radius:10px;
    background:#4e54c8;
    color:white;
    font-size:18px;
    font-weight:bold;
}

button:hover{
    background:#3b40a4;
}
</style>
</head>

<body>

<div class="card">
<div class="header" style="display:flex;justify-content:space-between;align-items:center;">
    <span>HIRWA SMART</span>
    <a href="/logout" style="color:white;text-decoration:none;font-size:14px;">Logout</a>
</div>
    <div class="subtitle">Login to continue</div>

    <form method="POST">
        <input type="text" name="username" placeholder="Username" required>
        <input type="password" name="password" placeholder="Password" required>
        <button type="submit">Login</button>
    </form>
</div>

</body>
</html>
"""
@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")
# ================= DASHBOARD =================
@app.route("/dashboard")
def dashboard():
    return """
<!DOCTYPE html>
<html>
<head>
<title>HIRWA SMART Dashboard</title>
<meta name="viewport" content="width=device-width, initial-scale=1">

<style>
body{
    margin:0;
    font-family: Arial, Helvetica, sans-serif;
    background:#f4f6fb;
}

.header{
    background:linear-gradient(90deg,#4e54c8,#8f94fb);
    color:white;
    padding:20px;
    text-align:center;
    font-size:22px;
    font-weight:bold;
}

.welcome{
    padding:20px;
    color:white;
}

.card{
    background:white;
    margin:15px;
    padding:18px;
    border-radius:15px;
    box-shadow:0 4px 10px rgba(0,0,0,0.08);
    display:flex;
    align-items:center;
    justify-content:space-between;
    text-decoration:none;
}

.card h2{
    margin:0;
}

.income{ border-left:8px solid #28a745; }
.expense{ border-left:8px solid #dc3545; }
.activity{ border-left:8px solid #007bff; }

.summary{
    margin:15px;
    background:white;
    padding:15px;
    border-radius:15px;
    box-shadow:0 4px 10px rgba(0,0,0,0.08);
}

.summary-box{
    display:flex;
    justify-content:space-between;
    margin-top:10px;
}

.box{
    width:32%;
    padding:12px;
    border-radius:10px;
    color:white;
    text-align:center;
    font-weight:bold;
}

.income-box{ background:#28a745; }
.expense-box{ background:#dc3545; }
.balance-box{ background:#007bff; }

a{ color:black; text-decoration:none; }
</style>
</head>

<body>

<div class="header">HIRWA SMART</div>

<div class="card income" onclick="location.href='/income'">
    <h2>💰 Income</h2>
    <span>Track your income</span>
</div>

<div class="card expense" onclick="location.href='/expenses'">
    <h2>💸 Expenses</h2>
    <span>Track your expenses</span>
</div>

<div class="card activity" onclick="location.href='/activities'">
    <h2>📋 Activities</h2>
    <span>View activities</span>
</div>

<div class="summary">
    <h3>Summary</h3>
    <div class="summary-box">
        <div class="box income-box">Income<br>RWF 0</div>
        <div class="box expense-box">Expenses<br>RWF 0</div>
        <div class="box balance-box">Balance<br>RWF 0</div>
    </div>
</div>

</body>
</html>
"""

# ================= INCOME =================
from flask import Response
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl import Workbook
from docx import Document



@app.route('/income', methods=['GET', 'POST'])
def income():
    if 'user_id' not in session:
        return redirect('/login')

    db = get_db()
    user_id = session['user_id']
    cur = db.cursor(dictionary=True)

    # DELETE (soft)
    delete_id = request.args.get('delete')
    if delete_id:
        cur.execute("""
            UPDATE income
            SET deleted_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (delete_id, user_id))
        db.commit()
        return redirect('/income')

    # EDIT LOAD
    edit_id = request.args.get('edit')
    edit_row = None
    if edit_id:
        cur.execute("SELECT * FROM income WHERE id=%s", (edit_id,))
        edit_row = cur.fetchone()

    # UPDATE
    if edit_id and request.method == 'POST':
        cur.execute("""
            UPDATE income
            SET amount=%s, source=%s, date=%s, note=%s,
                updated_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (
            request.form['amount'],
            request.form['source'],
            request.form['date'],
            request.form['note'],
            edit_id,
            user_id
        ))
        db.commit()
        return redirect('/income')

    # INSERT
    if request.method == 'POST':
        cur.execute("""
            INSERT INTO income(amount,source,date,note,user_id,user,created_at)
            VALUES(%s,%s,%s,%s,%s,%s,NOW())
        """, (
            request.form['amount'],
            request.form['source'],
            request.form['date'],
            request.form['note'],
            user_id,
            "user_"+str(user_id)
        ))
        db.commit()
        return redirect('/income')

    # SELECT (no deleted)
    cur.execute("""
        SELECT * FROM income
        WHERE user_id=%s AND deleted_at IS NULL
        ORDER BY id DESC
    """, (user_id,))
    rows = cur.fetchall()

    table = ""
    for r in rows:
        table += f"""
        <tr>
            <td>{r['amount']}</td>
            <td>{r['source']}</td>
            <td>{r['date']}</td>
            <td>{r['note']}</td>
            <td>
                <a href="/income?edit={r['id']}">Edit</a> |
                <a href="/income?delete={r['id']}">Delete</a>
            </td>
        </tr>
        """

    return f"""
    <h2>Income</h2>

    <form method="POST">
        Amount: <input name="amount" value="{edit_row['amount'] if edit_row else ''}"><br>
        Source: <input name="source" value="{edit_row['source'] if edit_row else ''}"><br>
        Date: <input type="date" name="date" value="{edit_row['date'] if edit_row else ''}"><br>
        Note: <input name="note" value="{edit_row['note'] if edit_row else ''}"><br>
        <button type="submit">Save</button>
    </form>

    <table border="1">
        <tr>
            <th>Amount</th>
            <th>Source</th>
            <th>Date</th>
            <th>Note</th>
            <th>Action</th>
        </tr>
        {table}
    </table>
    """
    # ===== FETCH =====
    cur.execute("SELECT * FROM income WHERE user_id=%s ORDER BY id DESC", (user_id,))
    rows = cur.fetchall()

    total = sum(float(r['amount']) for r in rows)

    # ===== HTML =====
    html = f"""
    <h2>Add Income</h2>
    <form method="POST">
    Amount:<input name="amount"><br>
    Source:<input name="source"><br>
    Date:<input type="date" name="date"><br>
    Note:<input name="note"><br>
    <button>Save</button>
    </form>

    <h3>Total Income: {total}</h3>

    <table border="1" cellpadding="8">
    <tr>
        <th>#</th><th>Amount</th><th>Source</th><th>Date</th><th>Note</th><th>Actions</th>
    </tr>
    """

    for i, r in enumerate(rows, 1):
        html += f"""
        <tr>
            <td>{i}</td>
            <td>{r['amount']}</td>
            <td>{r['source']}</td>
            <td>{r['date']}</td>
            <td>{r['note']}</td>
            <td>
                <a href="/income?delete={r['id']}">Delete</a>
            </td>
        </tr>
        """

    html += """
    </table>
    <br>
    <a href="/income/report/pdf">Download PDF</a> |
    <a href="/income/report/docx">Download DOC</a> |
    <a href="/income/report/excel">Download Excel</a>
    <br><br>
    <a href="/dashboard">Back</a>
    """
    return html
@app.route('/income/report/pdf')
def income_pdf():
    if 'user_id' not in session:
        return redirect('/login')

    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT * FROM income WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()

    file = "income_report.pdf"
    doc = SimpleDocTemplate(file)
    styles = getSampleStyleSheet()

    content = "INCOME REPORT\n\n"
    for r in rows:
        content += f"{r['amount']} - {r['source']} - {r['date']}\n"

    doc.build([Paragraph(content, styles['Normal'])])
    return send_file(file, as_attachment=True)


# ================= INCOME =================
@app.route('/income', methods=['GET', 'POST'])
def income_page():
    if 'user_id' not in session:
        return redirect('/login')

    db = get_db()
    cur = db.cursor()
    user_id = session['user_id']

    # DELETE
    delete_id = request.args.get('delete')
    if delete_id:
        cur.execute("UPDATE income SET deleted_at=NOW() WHERE id=%s AND user_id=%s", (delete_id, user_id))
        db.commit()
        return redirect('/income')

    # EDIT LOAD
    edit_id = request.args.get('edit')
    edit_row = None
    if edit_id:
        cur.execute("SELECT * FROM income WHERE id=%s AND user_id=%s", (edit_id, user_id))
        edit_row = cur.fetchone()

    # UPDATE
    if edit_id and request.method == 'POST':
        cur.execute("""
            UPDATE income
            SET amount=%s, source=%s, date=%s, note=%s, updated_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (
            request.form['amount'],
            request.form['source'],
            request.form['date'],
            request.form['note'],
            edit_id,
            user_id
        ))
        db.commit()
        return redirect('/income')

    # INSERT
    if request.method == 'POST' and not edit_id:
        cur.execute("""
            INSERT INTO income(amount,source,date,note,user_id,user,created_at)
            VALUES(%s,%s,%s,%s,%s,%s,NOW())
        """, (
            request.form['amount'],
            request.form['source'],
            request.form['date'],
            request.form['note'],
            user_id,
            "user_" + str(user_id)
        ))
        db.commit()
        return redirect('/income')

    # SELECT
    cur.execute("""
        SELECT * FROM income
        WHERE user_id=%s AND deleted_at IS NULL
        ORDER BY id DESC
    """, (user_id,))
    rows = cur.fetchall()

    total = sum(float(r[1]) for r in rows) if rows else 0

    table = ""
    for r in rows:
        table += f"""
        <tr>
            <td>{r[1]}</td>
            <td>{r[2]}</td>
            <td>{r[3]}</td>
            <td>{r[4]}</td>
            <td>
                <a href="/income?edit={r[0]}">Edit</a> |
                <a href="/income?delete={r[0]}">Delete</a>
            </td>
        </tr>
        """

    html = f"""
    <h2>Income</h2>

    <form method="POST">
        Amount: <input name="amount" value="{edit_row[1] if edit_row else ''}"><br>
        Source: <input name="source" value="{edit_row[2] if edit_row else ''}"><br>
        Date: <input type="date" name="date" value="{edit_row[3] if edit_row else ''}"><br>
        Note: <input name="note" value="{edit_row[4] if edit_row else ''}"><br>
        <button type="submit">Save</button>
    </form>

    <h3>Total Income: {total}</h3>

    <table border="1">
        <tr>
            <th>Amount</th><th>Source</th><th>Date</th><th>Note</th><th>Action</th>
        </tr>
        {table}
    </table>

    <br>
    <a href="/income/report/pdf">PDF</a> |
    <a href="/income/report/docx">DOC</a> |
    <a href="/income/report/excel">Excel</a>
    <br><br>
    <a href="/dashboard">Back</a>
    """

    cur.close()
    db.close()
    return html


# ================= PDF =================
@app.route('/income/report/pdf')
def income_pdf_report():
    if 'user_id' not in session:
        return redirect('/login')

    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT * FROM income WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()

    file = "income_report.pdf"
    doc = SimpleDocTemplate(file)
    styles = getSampleStyleSheet()

    content = "INCOME REPORT\n\n"
    for r in rows:
        content += f"{r[1]} - {r[2]} - {r[3]}\n"

    doc.build([Paragraph(content, styles['Normal'])])

    cur.close()
    db.close()
    return send_file(file, as_attachment=True)


# ================= DOCX =================
@app.route('/income/report/docx')
def income_docx():
    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT * FROM income WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()

    file = "income_report.docx"
    doc = Document()
    doc.add_heading("Income Report", 0)

    for r in rows:
        doc.add_paragraph(f"{r[1]} - {r[2]} - {r[3]}")

    doc.save(file)

    cur.close()
    db.close()
    return send_file(file, as_attachment=True)


# ================= EXCEL =================
@app.route('/income/report/excel')
def income_excel():
    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT * FROM income WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()

    file = "income_report.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["Amount", "Source", "Date", "Note"])

    for r in rows:
        ws.append([r[1], r[2], r[3], r[4]])

    wb.save(file)

    cur.close()
    db.close()
    return send_file(file, as_attachment=True)
# ================= EXPENSES =================

@app.route('/expenses', methods=['GET', 'POST'])
def expenses():
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    conn = get_db()
    cur = conn.cursor(pymysql.cursors.DictCursor)

    # DELETE (soft)
    delete_id = request.args.get('delete')
    if delete_id:
        cur.execute("""
            UPDATE expenses
            SET deleted_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (delete_id, user_id))
        conn.commit()
        conn.close()
        return redirect('/expenses')

    # EDIT FETCH
    edit_id = request.args.get('edit')
    edit_row = None
    if edit_id:
        cur.execute("SELECT * FROM expenses WHERE id=%s AND user_id=%s", (edit_id, user_id))
        edit_row = cur.fetchone()

    # UPDATE
    if edit_id and request.method == 'POST':
        cur.execute("""
            UPDATE expenses
            SET amount=%s, category=%s, date=%s, note=%s, updated_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (
            request.form['amount'],
            request.form['category'],
            request.form['date'],
            request.form['note'],
            edit_id,
            user_id
        ))
        conn.commit()
        conn.close()
        return redirect('/expenses')

    # INSERT
    if request.method == 'POST':
        cur.execute("""
            INSERT INTO expenses(amount,category,date,note,user_id,user,created_at)
            VALUES(%s,%s,%s,%s,%s,%s,NOW())
        """, (
            request.form['amount'],
            request.form['category'],
            request.form['date'],
            request.form['note'],
            user_id,
            "user_" + str(user_id)
        ))
        conn.commit()
        conn.close()
        return redirect('/expenses')

    # FETCH
    cur.execute("""
        SELECT * FROM expenses
        WHERE user_id=%s AND deleted_at IS NULL
        ORDER BY id DESC
    """, (user_id,))
    rows = cur.fetchall()
    conn.close()

    total = sum(float(r['amount']) for r in rows)

    table = ""
    for r in rows:
        table += f"""
        <tr>
            <td>{r['amount']}</td>
            <td>{r['category']}</td>
            <td>{r['date']}</td>
            <td>{r['note']}</td>
            <td>
                <a href="/expenses?edit={r['id']}">Edit</a> |
                <a href="/expenses?delete={r['id']}">Delete</a>
            </td>
        </tr>
        """

    return f"""
    <h2>Expenses</h2>

    <form method="POST">
        Amount: <input name="amount" value="{edit_row['amount'] if edit_row else ''}"><br>
        Category: <input name="category" value="{edit_row['category'] if edit_row else ''}"><br>
        Date: <input type="date" name="date" value="{edit_row['date'] if edit_row else ''}"><br>
        Note: <input name="note" value="{edit_row['note'] if edit_row else ''}"><br>
        <button type="submit">Save</button>
    </form>

    <h3>Total Expenses: {total}</h3>

    <table border="1" cellpadding="8">
        <tr>
            <th>Amount</th>
            <th>Category</th>
            <th>Date</th>
            <th>Note</th>
            <th>Action</th>
        </tr>
        {table}
    </table>

    <br>
    <a href="/expenses/report/pdf">Download PDF</a> |
    <a href="/expenses/report/docx">Download DOC</a> |
    <a href="/expenses/report/excel">Download Excel</a>
    <br><br>
    <a href="/dashboard">Back</a>
    """


@app.route('/expenses/report/pdf')
def expenses_pdf():
    conn = get_db()
    cur = conn.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT * FROM expenses WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()
    conn.close()

    file = "expenses_report.pdf"
    doc = SimpleDocTemplate(file)
    styles = getSampleStyleSheet()

    content = "EXPENSE REPORT\n\n"
    for r in rows:
        content += f"{r['amount']} - {r['category']} - {r['date']}\n"

    doc.build([Paragraph(content, styles['Normal'])])
    return send_file(file, as_attachment=True)


@app.route('/expenses/report/docx')
def expenses_docx():
    conn = get_db()
    cur = conn.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT * FROM expenses WHERE user_id=%s", (session['user_id'],))
    rows = cur.fetchall()
    conn.close()

    file = "expenses_report.docx"
    doc = Document()
    doc.add_heading("Expenses Report", 0)

    for r in rows:
        doc.add_paragraph(f"{r['amount']} - {r['category']} - {r['date']}")

    doc.save(file)
    return send_file(file, as_attachment=True)
@app.route('/expenses/report/excel')
def expenses_excel():
    if 'user_id' not in session:
        return redirect('/login')

    conn = get_db()
    cur = conn.cursor(pymysql.cursors.DictCursor)

    cur.execute(
        "SELECT amount, category, date, note FROM expenses WHERE user_id=%s AND deleted_at IS NULL ORDER BY id DESC",
        (session['user_id'],)
    )
    rows = cur.fetchall()
    conn.close()

    file = "expenses_report.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Expenses"

    # Header
    ws.append(["Amount", "Category", "Date", "Note"])

    # Data
    for r in rows:
        ws.append([r['amount'], r['category'], r['date'], r['note']])

    wb.save(file)
    return send_file(file, as_attachment=True)

# ================= ACTIVITIES =================
# ================= ACTIVITIES =================
@app.route("/activities", methods=["GET", "POST"])
def activities():
    conn = get_db()
    cur = conn.cursor()

    user_id = session.get("user_id")

    # 🔐 must be logged in (optional protection)
    if not user_id:
        return redirect("/login")

    # ================= ADD ACTIVITY =================
    if request.method == "POST" and "edit_id" not in request.form:
        activity_name = request.form["name"]
        done_by = request.form["done_by"]
        date = request.form["date"]
        desc = request.form["description"]

        cur.execute("""
            INSERT INTO activities
            (activity_name, done_by, date, description, user_id, created_at)
            VALUES (%s,%s,%s,%s,%s,NOW())
        """, (activity_name, done_by, date, desc, user_id))

        conn.commit()
        return redirect("/activities")

    # ================= EDIT LOAD =================
    edit_id = request.args.get("edit")
    edit_data = None

    if edit_id:
        cur.execute("SELECT * FROM activities WHERE id=%s", (edit_id,))
        edit_data = cur.fetchone()

    # ================= UPDATE =================
    if request.method == "POST" and "edit_id" in request.form:
        cur.execute("""
            UPDATE activities
            SET activity_name=%s,
                done_by=%s,
                date=%s,
                description=%s,
                updated_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (
            request.form["name"],
            request.form["done_by"],
            request.form["date"],
            request.form["description"],
            request.form["edit_id"],
            user_id
        ))
        conn.commit()
        return redirect("/activities")

    # ================= DELETE (SOFT = REJECTED) =================
    delete_id = request.args.get("delete")
    if delete_id:
        cur.execute("""
            UPDATE activities
            SET deleted_at=NOW()
            WHERE id=%s AND user_id=%s
        """, (delete_id, user_id))
        conn.commit()
        return redirect("/activities")

    # ================= APPROVE / REJECT =================
    approve_id = request.args.get("approve")
    reject_id = request.args.get("reject")

    if approve_id:
        cur.execute("""
            UPDATE activities SET status='approved'
            WHERE id=%s AND user_id=%s
        """, (approve_id, user_id))
        conn.commit()
        return redirect("/activities")

    if reject_id:
        cur.execute("""
            UPDATE activities SET status='rejected'
            WHERE id=%s AND user_id=%s
        """, (reject_id, user_id))
        conn.commit()
        return redirect("/activities")

    # ================= FETCH =================
    cur.execute("""
        SELECT * FROM activities
        WHERE user_id=%s AND deleted_at IS NULL
        ORDER BY id DESC
    """, (user_id,))
    rows = cur.fetchall()

    html_rows = ""
    for r in rows:
        html_rows += f"""
        <div class='activity-item'>
            <b>{r['activity_name']}</b>
            <small>({r.get('status','pending')})</small><br>

            👤 {r['done_by']} | 📅 {r['date']}<br>
            📝 {r['description']}<br>

            <a href="/activities?edit={r['id']}">Edit</a> |
            <a href="/activities?delete={r['id']}">Delete</a> |
            <a href="/activities?approve={r['id']}">Approve</a> |
            <a href="/activities?reject={r['id']}">Reject</a>
        </div>
        """

    return f"""
<!DOCTYPE html>
<html>
<head>
<title>Activities - HIRWA SMART</title>
<meta name="viewport" content="width=device-width, initial-scale=1">

<style>
body{{margin:0;font-family:Arial;background:#f4f6fb;}}
.header{{background:linear-gradient(90deg,#4e54c8,#8f94fb);color:white;padding:18px;text-align:center;font-size:20px;font-weight:bold;}}
.container{{padding:15px;}}
.card{{background:white;border-radius:15px;padding:15px;box-shadow:0 4px 12px rgba(0,0,0,0.08);margin-bottom:15px;}}
input,textarea{{width:100%;padding:12px;margin:8px 0;border-radius:10px;border:1px solid #ddd;}}
button{{width:100%;padding:14px;border:none;border-radius:10px;background:#4e54c8;color:white;font-weight:bold;}}
.activity-item{{background:#eef1ff;padding:12px;border-radius:10px;margin-bottom:10px;}}
</style>
</head>

<body>

<div class="header">HIRWA SMART - Activities</div>

<div class="container">

<div class="card">
<h3>Add / Edit Activity</h3>

<form method="POST">
<input name="name" placeholder="Activity name"
value="{edit_data['activity_name'] if edit_data else ''}" required>

<input name="done_by" placeholder="Done by"
value="{edit_data['done_by'] if edit_data else ''}" required>

<input type="date" name="date"
value="{edit_data['date'] if edit_data else ''}" required>

<textarea name="description" placeholder="Description">{edit_data['description'] if edit_data else ''}</textarea>

{"<input type='hidden' name='edit_id' value='"+str(edit_data['id'])+"'>" if edit_data else ""}

<button type="submit">Save</button>
</form>
</div>

<div class="card">
<h3>All Activities</h3>
{html_rows}
</div>

</div>

</body>
</html>
"""
@app.route("/fixdb")
def fix_db():
    db = get_db()
    cur = db.cursor()

    cur.execute("""
    ALTER TABLE income
    ADD COLUMN user VARCHAR(100),
    ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    ADD COLUMN updated_at DATETIME NULL,
    ADD COLUMN deleted_at DATETIME NULL;
    """)

    cur.execute("""
    ALTER TABLE expenses
    ADD COLUMN user VARCHAR(100),
    ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    ADD COLUMN updated_at DATETIME NULL,
    ADD COLUMN deleted_at DATETIME NULL;
    """)

    cur.execute("""
    ALTER TABLE activities
    ADD COLUMN user VARCHAR(100),
    ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    ADD COLUMN updated_at DATETIME NULL,
    ADD COLUMN deleted_at DATETIME NULL;
    """)

    db.commit()
    return "DB FIXED ✅"
@app.errorhandler(Exception)
def handle_error(e):
    return f"""
    <h2>SERVER ERROR</h2>
    <pre>{e}</pre>
    """, 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
